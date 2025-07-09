from docx import Document

def replace_placeholders(doc, replacements):
    def replace_in_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ''

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

def fill_template_to_bytes(template_path, data):
    from io import BytesIO
    doc = Document(template_path)
    replace_placeholders(doc, data)
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f
